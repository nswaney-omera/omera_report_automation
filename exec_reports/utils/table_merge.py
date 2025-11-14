"""
Utility for merging duplicate cells in Word document tables.

This module handles merging consecutive rows with the same value
in the first column of a table (commonly used for employment history).
"""

from docx import Document
from pathlib import Path
import sys
import io


def check_mergeable_cells(table):
    """
    Check which cells in the first column can be merged.

    Args:
        table: A docx table object

    Returns:
        List of tuples (start_row, end_row, value) indicating mergeable cell ranges
    """
    mergeable_ranges = []
    current_value = None
    start_row = None

    for row_idx, row in enumerate(table.rows):
        if len(row.cells) == 0:
            continue

        cell_value = row.cells[0].text.strip()

        # Skip empty cells
        if not cell_value:
            if current_value is not None:
                # End the current range
                if start_row is not None and row_idx - 1 > start_row:
                    mergeable_ranges.append((start_row, row_idx - 1, current_value))
                current_value = None
                start_row = None
            continue

        # Check if this is the same as the previous value
        if cell_value == current_value:
            # Continue the current range
            continue
        else:
            # Value changed - save the previous range if it was more than 1 row
            if current_value is not None and start_row is not None:
                if row_idx - 1 > start_row:
                    mergeable_ranges.append((start_row, row_idx - 1, current_value))

            # Start a new range
            current_value = cell_value
            start_row = row_idx

    # Don't forget the last range
    if current_value is not None and start_row is not None:
        if len(table.rows) - 1 > start_row:
            mergeable_ranges.append((start_row, len(table.rows) - 1, current_value))

    return mergeable_ranges


def merge_cells_in_column(table, start_row, end_row, column=0):
    """
    Merge cells vertically in a given column from start_row to end_row.
    After merging, keeps only one copy of the text (removes duplicates).

    Args:
        table: A docx table object
        start_row: Starting row index (0-based)
        end_row: Ending row index (0-based, inclusive)
        column: Column index to merge (default: 0, first column)

    Returns:
        bool: True if merge was successful, False otherwise
    """
    try:
        # Validation: Check that user input is correct
        if start_row >= end_row:
            print(f"Warning: start_row ({start_row}) must be less than end_row ({end_row})")
            return False

        if start_row < 0 or end_row >= len(table.rows):
            print(f"Warning: Invalid row indices ({start_row}, {end_row}) for table with {len(table.rows)} rows")
            return False

        # Get the cells to merge
        start_cell = table.rows[start_row].cells[column]
        end_cell = table.rows[end_row].cells[column]

        # Store the original text value (before merge duplicates it)
        original_text = start_cell.text.strip()

        # Perform the merge
        start_cell.merge(end_cell)

        # Clear all content from the merged cell
        for paragraph in start_cell.paragraphs:
            paragraph.clear()

        # Set the text to only one copy of the original value
        start_cell.text = original_text

        return True

    except Exception as e:
        print(f"Error merging cells: {str(e)}")
        return False


def analyze_table_for_merge(table, table_idx):
    """
    Analyze a table and report which cells can be merged.

    Args:
        table: A docx table object
        table_idx: Index of the table in the document (for reporting)

    Returns:
        Dict with analysis results
    """
    mergeable = check_mergeable_cells(table)

    result = {
        "table_index": table_idx,
        "total_rows": len(table.rows),
        "total_columns": len(table.rows[0].cells) if table.rows else 0,
        "mergeable_ranges": len(mergeable),
        "ranges": mergeable
    }

    return result


def merge_duplicate_first_column_cells(input_path, output_path, table_index=None):
    """
    Main function: Detect and merge duplicate cells in the first column of a specific table.

    This function:
    1. Loads a Word document
    2. Checks the specified table for consecutive rows with the same value in column 0
    3. Merges those cells vertically
    4. Saves the modified document

    Args:
        input_path: Path to the input Word document
        output_path: Path where the modified document will be saved
        table_index: Index of the table to process (0-based). If None, processes all tables.

    Returns:
        Dict with success status, message, and merge statistics
    """
    try:
        # Load the document
        doc = Document(input_path)

        if len(doc.tables) == 0:
            return {
                "success": False,
                "message": "No tables found in document"
            }

        # Validate table index if specified
        if table_index is not None:
            if table_index < 0 or table_index >= len(doc.tables):
                return {
                    "success": False,
                    "message": f"Invalid table index {table_index}. Document has {len(doc.tables)} tables (0-{len(doc.tables)-1})"
                }

        total_merges = 0
        table_results = []

        # Determine which tables to process
        if table_index is not None:
            tables_to_process = [(table_index, doc.tables[table_index])]
            print(f"\n=== Processing table {table_index} only ===\n")
        else:
            tables_to_process = list(enumerate(doc.tables))
            print(f"\n=== Analyzing {len(doc.tables)} tables ===\n")

        # Analyze specified table(s)
        for table_idx, table in tables_to_process:
            analysis = analyze_table_for_merge(table, table_idx)
            table_results.append(analysis)

            print(f"Table {table_idx}:")
            print(f"  Rows: {analysis['total_rows']}, Columns: {analysis['total_columns']}")
            print(f"  Mergeable ranges: {analysis['mergeable_ranges']}")

            if analysis['mergeable_ranges'] > 0:
                print(f"  Ranges to merge:")
                for start, end, value in analysis['ranges']:
                    print(f"    Rows {start}-{end}: '{value}'")
                    # Perform the merge
                    if merge_cells_in_column(table, start, end, column=0):
                        total_merges += 1
                        print(f"    [OK] Merged successfully")
                    else:
                        print(f"    [FAIL] Merge failed")
            print()

        # Save the modified document
        doc.save(output_path)

        if table_index is not None:
            message = f"Successfully processed table {table_index} and performed {total_merges} merges"
        else:
            message = f"Successfully analyzed {len(doc.tables)} tables and performed {total_merges} merges"

        return {
            "success": True,
            "message": message,
            "input_file": input_path,
            "output_file": output_path,
            "table_index": table_index,
            "total_merges": total_merges,
            "table_results": table_results,
            "cells_merged": total_merges  # For compatibility with server.py
        }

    except Exception as e:
        # Handle encoding issues in error messages
        error_msg = str(e)
        try:
            # Try to convert to ASCII-safe string
            error_msg = error_msg.encode('ascii', 'replace').decode('ascii')
        except:
            error_msg = "Error processing document (encoding issue in error message)"
        
        return {
            "success": False,
            "message": f"Error processing document: {error_msg}"
        }


def test_table_merge(table_index=None):
    """
    Test function: Run the merge detection and merging on the test document.

    Args:
        table_index: Optional index of specific table to process (0-based).
                    If None, processes all tables.
    """
    input_file = "Simone_Tregeagle_20251027_140617_report_20251027_140626.docx"
    output_file = "Simone_Tregeagle_merged_output.docx"

    print("=== Table Merge Test Script ===")
    print(f"Input: {input_file}")
    print(f"Output: {output_file}")
    if table_index is not None:
        print(f"Target table: {table_index}")
    else:
        print(f"Target table: All tables")

    result = merge_duplicate_first_column_cells(input_file, output_file, table_index=table_index)

    print("\n=== Results ===")
    print(f"Success: {result['success']}")
    print(f"Message: {result['message']}")

    if result['success']:
        print(f"\nStatistics:")
        print(f"  Total merges performed: {result['total_merges']}")

        print(f"\nDetailed results:")
        for table_result in result['table_results']:
            if table_result['mergeable_ranges'] > 0:
                print(f"  Table {table_result['table_index']}: {table_result['mergeable_ranges']} merge(s)")

    return result


if __name__ == "__main__":
    # Run the test
    # Example: test_table_merge(table_index=0)  # Process only table 0
    # Example: test_table_merge()                # Process all tables
    test_table_merge(table_index=0)  # By default, only process table 0