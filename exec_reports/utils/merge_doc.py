"""
Utility for merging multiple Word documents into a single shortlist report.

This module combines:
- Overview document (summary, matrix, candidate list)
- Individual candidate profile documents
- End page document

All sections are assembled with proper page breaks.
"""

from docxcompose.composer import Composer
from docx import Document
from pathlib import Path


def merge_shortlist_documents(
    overview_path: str,
    candidate_paths: list,
    output_path: str,
    end_page_path: str
) -> dict:
    """
    Merge overview, candidate profiles, and end page into a single shortlist report.

    The documents are assembled in this order:
    1. Overview section (summary, matrix, candidate list)
    2. Individual candidate documents (in provided order)
    3. End page

    Each section starts on a new page.

    Args:
        overview_path: Path to overview document (6 pages)
        candidate_paths: List of paths to candidate documents (3 pages each, in order)
        output_path: Path where merged document will be saved
        end_page_path: Path to end page document

    Returns:
        Dictionary with:
        - success (bool): Whether merge succeeded
        - message (str): Success/error message
        - output_file (str): Path to merged document (if successful)
        - sections_merged (int): Number of sections combined (if successful)

    Example:
        result = merge_shortlist_documents(
            overview_path="overview.docx",
            candidate_paths=["candidate1.docx", "candidate2.docx"],
            output_path="final_report.docx",
            end_page_path="end_page.docx"
        )
    """
    try:
        # Validate inputs
        overview_file = Path(overview_path)
        if not overview_file.exists():
            return {
                "success": False,
                "message": f"Overview document not found: {overview_path}",
                "output_file": None
            }

        end_page_file = Path(end_page_path)
        if not end_page_file.exists():
            return {
                "success": False,
                "message": f"End page document not found: {end_page_path}",
                "output_file": None
            }

        # Validate candidate documents
        for i, candidate_path in enumerate(candidate_paths, 1):
            candidate_file = Path(candidate_path)
            if not candidate_file.exists():
                return {
                    "success": False,
                    "message": f"Candidate document {i} not found: {candidate_path}",
                    "output_file": None
                }

        # Start with overview document
        master = Document(overview_path)
        composer = Composer(master)

        # Append each candidate document
        for candidate_path in candidate_paths:
            candidate_doc = Document(candidate_path)
            composer.append(candidate_doc)

        # Append end page
        end_doc = Document(end_page_path)
        composer.append(end_doc)

        # Save merged document
        output_file = Path(output_path)
        output_file.parent.mkdir(parents=True, exist_ok=True)
        composer.save(str(output_file))

        sections_merged = 1 + len(candidate_paths) + 1  # overview + candidates + end page

        return {
            "success": True,
            "message": f"Successfully merged {sections_merged} sections into report",
            "output_file": str(output_file),
            "sections_merged": sections_merged,
            "overview": str(overview_file.name),
            "candidates": len(candidate_paths),
            "end_page": str(end_page_file.name)
        }

    except Exception as e:
        return {
            "success": False,
            "message": f"Error merging documents: {str(e)}",
            "output_file": None
        }


def merge_documents(start_path: str, candidate_paths: list, end_path: str, output_path: str = None) -> Path:
    """
    Legacy function for backwards compatibility.
    Wraps merge_shortlist_documents with simpler interface.

    Args:
        start_path: Path to overview/start document
        candidate_paths: List of candidate document paths
        end_path: Path to end page document
        output_path: Optional output path (defaults to 'combined.docx' in current directory)

    Returns:
        Path to merged document

    Raises:
        Exception if merge fails
    """
    if output_path is None:
        output_path = Path.cwd() / "combined.docx"

    result = merge_shortlist_documents(
        overview_path=start_path,
        candidate_paths=candidate_paths,
        output_path=str(output_path),
        end_page_path=end_path
    )

    if not result["success"]:
        raise Exception(result["message"])

    return Path(result["output_file"])

if __name__ == "__main__":
    # Test the document merger
    print("Testing document merger...")
    
    result = merge_shortlist_documents(
        overview_path=r"src\exec_reports\outputs\Chief_Technology_Officer_Quantium_2025\merged_candidates_Chief_Technology_Officer_Quantium_2025_report_20251105_175609.docx",
        candidate_paths=[
            r"src\exec_reports\outputs\Chief_Technology_Officer_Quantium_2025\temp_Eleanor_Debelle_20251105_175505_report_20251105_175610.docx",
            r"src\exec_reports\outputs\Chief_Technology_Officer_Quantium_2025\temp_Justin_James_20251105_175307_report_20251105_175610.docx",
            r"src\exec_reports\outputs\Chief_Technology_Officer_Quantium_2025\temp_Sarah_Morris_20251105_174743_report_20251105_175610.docx"
        ],
        output_path="test_merged_report.docx",
        end_page_path=r"src\exec_reports\templates\s_end_page.docx"
    )
    
    # Print results
    if result["success"]:
        print(f"✓ {result['message']}")
        print(f"  Output file: {result['output_file']}")
        print(f"  Sections merged: {result['sections_merged']}")
        print(f"  Overview: {result['overview']}")
        print(f"  Candidates: {result['candidates']}")
        print(f"  End page: {result['end_page']}")
    else:
        print(f"✗ Error: {result['message']}")
