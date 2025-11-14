from docx import Document
from docx.oxml import parse_xml
from docx.oxml.ns import nsdecls

def set_cell_color(cell, color_hex):
    """Set cell background color"""
    shading_elm = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{color_hex}"/>')
    cell._element.get_or_add_tcPr().append(shading_elm)

def color_cell_by_rating(cell, rating):
    """Change cell color based on rating value"""
    try:
        rating = int(rating)

        if 1 <= rating <= 2:
            set_cell_color(cell, 'FADBB7')  # Very Low - light tan
        elif 3 <= rating <= 4:
            set_cell_color(cell, 'D1B399')  # Less Than Required - tan
        elif 5 <= rating <= 6:
            set_cell_color(cell, 'BB8E66')  # At Required Level - gold
        elif 7 <= rating <= 8:
            set_cell_color(cell, 'F2BA6E')  # Above Required Level - orange
        elif 9 <= rating <= 10:
            set_cell_color(cell, 'F97012')  # Very High - bright orange
    except:
        pass  # Skip if not a valid number

def get_color_by_rating(rating):
    """Get color hex based on rating value"""
    try:
        rating = int(rating)
        if 1 <= rating <= 2:
            return 'FADBB7'
        elif 3 <= rating <= 4:
            return 'D1B399'
        elif 5 <= rating <= 6:
            return 'BB8E66'
        elif 7 <= rating <= 8:
            return 'F2BA6E'
        elif 9 <= rating <= 10:
            return 'F97012'
    except:
        return None

def extract_ratings_from_summary_table(table):
    """
    Extract candidate ratings from table 4 (summary table)
    Returns a list of ratings for each candidate (7 skills each)
    """
    ratings_list = []

    # Get candidate rows (skip header row)
    for row_idx, row in enumerate(table.rows):
        if row_idx == 0:  # Skip header row
            continue

        candidate_name = row.cells[0].text.strip()
        if not candidate_name:
            continue

        # Extract ratings for 7 skills (columns 1-7)
        candidate_ratings = []
        for col_idx in range(1, 8):  # Columns 1-7 contain the 7 skill ratings
            if col_idx < len(row.cells):
                rating = row.cells[col_idx].text.strip()
                try:
                    candidate_ratings.append(int(rating))
                except:
                    candidate_ratings.append(5)  # Default to 5 if invalid

        ratings_list.append({
            'name': candidate_name,
            'ratings': candidate_ratings
        })

    return ratings_list

def color_candidate_detail_table(table, candidate_ratings):
    """
    Color cells in rows 1-7, columns 3-12 based on ratings.
    The number of colored cells corresponds to the rating value.
    """
    # Rows 1-7 correspond to the 7 skills
    for skill_idx in range(7):
        row_idx = skill_idx + 1  # Row 1-7 (0-indexed, so actually 1-7)

        if row_idx >= len(table.rows):
            continue

        row = table.rows[row_idx]
        rating = candidate_ratings[skill_idx]
        color_hex = get_color_by_rating(rating)

        if color_hex:
            # Color cells based on rating (columns 3-12)
            # Rating 10 = color all 10 cells (cols 3-12)
            # Rating 8 = color 8 cells (cols 3-10)
            # Rating 5 = color 5 cells (cols 3-7)
            # etc.
            cells_to_color = min(rating, 10)  # Max 10 cells

            for col_idx in range(3, 3 + cells_to_color):  # Start at col 3
                if col_idx < len(row.cells):
                    set_cell_color(row.cells[col_idx], color_hex)


def apply_table_colors(input_path, output_path):
    """
    Apply color formatting to shortlist report tables.

    Process:
    1. Colors Table 4 (Summary Table) cells based on their rating values
    2. Extracts ratings from Table 4
    3. Applies bar chart coloring to Tables 5-14 (Candidate Detail Tables)

    Args:
        input_path: Path to the input Word document
        output_path: Path where the colored document will be saved

    Returns:
        Dict with success status and message
    """
    try:
        # Load the document
        doc = Document(input_path)

        # Step 1: Process Table 4 (Summary Table)
        if len(doc.tables) < 5:
            return {
                "success": False,
                "message": f"Document has insufficient tables. Expected at least 5, found {len(doc.tables)}"
            }

        summary_table = doc.tables[4]

        # Color cells based on their content in table 4
        for row in summary_table.rows:
            for cell in row.cells:
                color_cell_by_rating(cell, cell.text.strip())

        # Step 2: Extract ratings from Table 4
        ratings_list = extract_ratings_from_summary_table(summary_table)

        # Step 3: Process Tables 5-14 (Individual Candidate Tables)
        candidates_processed = 0
        for table_idx in range(5, 15):  # Tables 5-14
            if table_idx < len(doc.tables):
                table = doc.tables[table_idx]
                candidate_idx = table_idx - 5  # Table 5 = candidate 0, etc.

                if candidate_idx < len(ratings_list):
                    candidate_data = ratings_list[candidate_idx]
                    candidate_ratings = candidate_data['ratings']
                    color_candidate_detail_table(table, candidate_ratings)
                    candidates_processed += 1

        # Save the modified document
        doc.save(output_path)

        return {
            "success": True,
            "message": f"Successfully colored tables for {candidates_processed} candidates",
            "input_file": input_path,
            "output_file": output_path,
            "candidates_processed": candidates_processed
        }

    except Exception as e:
        return {
            "success": False,
            "message": f"Error applying table colors: {str(e)}"
        }
