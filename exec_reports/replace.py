import json
import os
from docx import Document
from datetime import datetime

def replace_text_in_runs(runs, replacements):
    """Helper function to replace text in a collection of runs, preserving formatting"""
    # First, try simple replacement within individual runs
    for run in runs:
        if run.text:
            original_text = run.text
            for placeholder, value in replacements.items():
                if placeholder in original_text:
                    run.text = original_text.replace(placeholder, value)
                    original_text = run.text

    # Handle placeholders that might be split across runs
    full_text = ''.join(run.text for run in runs)

    for placeholder, value in replacements.items():
        if placeholder in full_text:
            # Find which runs contain parts of the placeholder
            current_pos = 0
            placeholder_start = full_text.find(placeholder)

            if placeholder_start >= 0:
                placeholder_end = placeholder_start + len(placeholder)

                # Find the runs that contain the placeholder
                run_positions = []
                pos = 0
                for i, run in enumerate(runs):
                    run_start = pos
                    run_end = pos + len(run.text)
                    run_positions.append((i, run_start, run_end, run.text))
                    pos = run_end

                # Find which runs overlap with the placeholder
                affected_runs = []
                for i, start, end, text in run_positions:
                    if start < placeholder_end and end > placeholder_start:
                        affected_runs.append((i, start, end, text))

                if affected_runs:
                    # Replace the placeholder across the affected runs
                    for i, (run_idx, start, end, text) in enumerate(affected_runs):
                        run = runs[run_idx]

                        if i == 0:  # First affected run gets the replacement
                            # Calculate how much of the placeholder is in this run
                            placeholder_in_run_start = max(0, placeholder_start - start)
                            placeholder_in_run_end = min(len(text), placeholder_end - start)

                            new_text = (text[:placeholder_in_run_start] +
                                      value +
                                      text[placeholder_in_run_end:])
                            run.text = new_text
                        else:  # Subsequent runs get their placeholder parts removed
                            placeholder_in_run_start = max(0, placeholder_start - start)
                            placeholder_in_run_end = min(len(text), placeholder_end - start)

                            new_text = (text[:placeholder_in_run_start] +
                                      text[placeholder_in_run_end:])
                            run.text = new_text

                # Recalculate full_text for next placeholder
                full_text = ''.join(run.text for run in runs)

def replace_all_text(doc, replacements):
    """Replace text in all parts of the document"""
    # Replace in paragraphs
    for paragraph in doc.paragraphs:
        if paragraph.runs:  # Only process if there are runs
            replace_text_in_runs(paragraph.runs, replacements)

    # Replace in tables
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    if paragraph.runs:  # Only process if there are runs
                        replace_text_in_runs(paragraph.runs, replacements)

    # Replace in headers
    for section in doc.sections:
        # Replace in header
        header = section.header
        for paragraph in header.paragraphs:
            if paragraph.runs:  # Only process if there are runs
                replace_text_in_runs(paragraph.runs, replacements)

        # Replace in footer
        footer = section.footer
        for paragraph in footer.paragraphs:
            if paragraph.runs:  # Only process if there are runs
                replace_text_in_runs(paragraph.runs, replacements)

        # Replace in footer tables
        for table in footer.tables:
            for row in table.rows:
                for cell in row.cells:
                    for paragraph in cell.paragraphs:
                        if paragraph.runs:
                            replace_text_in_runs(paragraph.runs, replacements)

def process_document_template(candidate_data_path=None,
                            template_path=None,
                            output_dir='src/outputs'):
    """
    Main function to process document template with candidate data

    Args:
        candidate_data_path: Path to JSON file containing candidate data
        template_path: Path to Word template document
        output_dir: Directory to save output file

    Returns:
        dict: Result with success status, message, and output file path
    """
    try:
        # Check if files exist
        if not os.path.exists(candidate_data_path):
            return {
                "success": False,
                "message": f"Candidate data file not found: {candidate_data_path}",
                "output_file": None
            }

        if not os.path.exists(template_path):
            return {
                "success": False,
                "message": f"Template file not found: {template_path}",
                "output_file": None
            }

        # Create output directory if it doesn't exist
        os.makedirs(output_dir, exist_ok=True)

        # Load candidate data
        with open(candidate_data_path, 'r') as f:
            data = json.load(f)

        # Create replacements dictionary
        replacements = {f'{{{k}}}': str(v) for k, v in data.items()}

        # Load document
        doc = Document(template_path)

        # Perform replacements
        replace_all_text(doc, replacements)

        # Generate output filename with timestamp
        timestamp = datetime.now().strftime('%Y%m%d_%H%M%S')
        # Extract just the filename from the full path and remove extension
        input_filename = os.path.splitext(os.path.basename(candidate_data_path))[0]
        output_filename = f'{input_filename}_report_{timestamp}.docx'
        output_path = os.path.join(output_dir, output_filename)

        # Save document
        doc.save(output_path)

        return {
            "success": True,
            "message": f"Document processed successfully. Replaced {len(replacements)} placeholders.",
            "output_file": output_path
        }

    except json.JSONDecodeError as e:
        return {
            "success": False,
            "message": f"Error reading JSON file: {str(e)}",
            "output_file": None
        }
    except Exception as e:
        return {
            "success": False,
            "message": f"Error processing document: {str(e)}",
            "output_file": None
        }

# Keep the original functionality for direct script execution
if __name__ == "__main__":
    result = process_document_template(candidate_data_path=r"C:\Users\nswaney\Documents\repositories\cr_automation\src\exec_reports\data\Chief_Executive_Officer_AMP_Limited_2025\Jenny_Oliver_20251027_150239.json", template_path=r'src\exec_reports\templates\s_singular_candidate.docx', output_dir=r'src\shortlist_report\outputs')
    if result["success"]:
        print(f"Success: {result['message']}")
        print(f"Output file: {result['output_file']}")
    else:
        print(f"Error: {result['message']}")